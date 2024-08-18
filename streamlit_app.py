import streamlit as st
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from docx import Document
from io import BytesIO
import time
import base64

def setup_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    # Use default ChromeDriver for Streamlit Cloud
    driver = webdriver.Chrome(service=Service(), options=chrome_options)
    return driver

@st.cache_resource
def search_naver_news(keyword, num_articles):
    driver = setup_driver()
    url = f"https://search.naver.com/search.naver?where=news&query={keyword}&sm=tab_jum&sort=0&photo=0&field=0&reporter_article=&pd=0&ds=&de=&docid=&nso=so:r,p:all,a:all&mynews=0&cluster_rank=53&start=1&refresh_start=0"
    driver.get(url)

    news_list = []
    seen_titles = set()
    
    try:
        while len(news_list) < num_articles:
            WebDriverWait(driver, 10).until(
                EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.list_news > li"))
            )

            news_elements = driver.find_elements(By.CSS_SELECTOR, "ul.list_news > li")

            for element in news_elements:
                if len(news_list) >= num_articles:
                    break
                try:
                    title = element.find_element(By.CSS_SELECTOR, "a.news_tit").text
                    if title not in seen_titles:
                        seen_titles.add(title)
                        summary = element.find_element(By.CSS_SELECTOR, "div.news_dsc").text
                        link = element.find_element(By.CSS_SELECTOR, "a.news_tit").get_attribute("href")
                        news_list.append({
                            'title': title,
                            'summary': summary,
                            'link': link
                        })
                except Exception as e:
                    print(f"기사 데이터 추출 중 오류 발생: {e}")

            if len(news_list) < num_articles:
                try:
                    next_button = driver.find_element(By.CSS_SELECTOR, "a.btn_next")
                    next_button.click()
                    time.sleep(1)
                except:
                    break
            else:
                break

    except Exception as e:
        print(f"뉴스 검색 중 오류 발생: {e}")
    finally:
        driver.quit()

    return news_list[:num_articles]

def create_word_document(news_list, keyword):
    doc = Document()
    doc.add_heading(f"'{keyword}' 관련 뉴스", 0)

    for idx, news in enumerate(news_list, 1):
        doc.add_heading(f"{idx}. {news['title']}", level=1)
        doc.add_paragraph(news['summary'])
        doc.add_paragraph(f"링크: {news['link']}")
        doc.add_paragraph("\n")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_download_link(buffer, filename):
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">다운로드</a>'

st.set_page_config(page_title="네이버 뉴스 크롤러", page_icon="📰", layout="wide")

st.title("📰 네이버 뉴스 크롤러")

st.markdown("""
    <style>
    .stApp {
        max-width: 800px;
        margin: 0 auto;
    }
    .news-item {
        background-color: #f0f2f6;
        border-radius: 5px;
        padding: 10px;
        margin-bottom: 10px;
    }
    .news-title {
        font-weight: bold;
        color: #1e3a8a;
    }
    .news-summary {
        font-size: 0.9em;
        color: #374151;
    }
    .stButton > button {
        width: 100%;
    }
    </style>
""", unsafe_allow_html=True)

keyword = st.text_input("검색할 키워드를 입력하세요:")
num_articles = st.slider("가져올 기사 수", 1, 20, 5)

col1, col2, col3 = st.columns([2, 2, 1])

if col2.button("뉴스 검색"):
    if keyword:
        with st.spinner("뉴스를 검색 중입니다..."):
            news_results = search_naver_news(keyword, num_articles)
        
        if news_results:
            st.session_state.news_results = news_results
            st.success(f"'{keyword}'에 대한 상위 {len(news_results)}개 뉴스를 찾았습니다!")
            for idx, news in enumerate(news_results, 1):
                st.markdown(f"""
                <div class="news-item">
                    <p class="news-title">{idx}. {news['title']}</p>
                    <p class="news-summary">{news['summary']}</p>
                    <a href="{news['link']}" target="_blank">기사 읽기</a>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.warning("검색 결과가 없습니다. 다른 키워드로 시도해보세요.")
    else:
        st.error("키워드를 입력해주세요.")

if col3.button("다운로드"):
    if 'news_results' in st.session_state and st.session_state.news_results:
        word_buffer = create_word_document(st.session_state.news_results, keyword)
        st.markdown(get_download_link(word_buffer, f"{keyword}_news.docx"), unsafe_allow_html=True)
    else:
        st.warning("먼저 뉴스를 검색해주세요.")
